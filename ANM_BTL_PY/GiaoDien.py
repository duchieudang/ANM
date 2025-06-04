import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from docx import Document
import fitz  # PyMuPDF
import random
from dsa import hash_signature, calculate_g, generate_k, sign, verify, is_prime

class DSAGui:
    def __init__(self, root):
        self.root = root
        root.title("DSA Digital Signature Tool")
        root.geometry("1200x750")  # Set width to accommodate three equal-width columns

        self.setup_style()

        self.r = None
        self.s = None
        self.signature_hash = None
        self.message = None

        self.p_var = tk.StringVar()
        self.q_var = tk.StringVar()
        self.h_var = tk.StringVar()
        self.g_var = tk.StringVar()
        self.x_var = tk.StringVar()
        self.y_var = tk.StringVar()

        self.setup_widgets()

    def setup_style(self):
        style = ttk.Style()
        style.theme_use("clam")
        # General styles
        style.configure("TLabel", font=('Segoe UI', 10), foreground="#333")
        style.configure("TEntry", font=('Segoe UI', 10), fieldbackground="#fff")
        style.configure("TFrame", background="#f5f5dc")
        style.configure("TLabelframe.Label", font=('Segoe UI', 11, 'bold'), foreground="#333")

        # Frame styles with pastel colors
        style.configure("Param.TLabelframe", background="#e6e6fa")  # Light lavender
        style.configure("Sign.TLabelframe", background="#d4edda")  # Light mint
        style.configure("Verify.TLabelframe", background="#ffe5d9")  # Light peach
        style.configure("ButtonFrame.TFrame", background="#f5f5dc")

        # Button styles with pastel colors and hover effects
        style.configure("Word.TButton", padding=6, font=('Segoe UI', 10), background="#d1ecf1", foreground="#333")  # Soft blue
        style.map("Word.TButton", background=[("active", "#bee5eb")])
        style.configure("PDF.TButton", padding=6, font=('Segoe UI', 10), background="#ffecd9", foreground="#333")  # Soft peach
        style.map("PDF.TButton", background=[("active", "#ffdab9")])
        style.configure("Excel.TButton", padding=6, font=('Segoe UI', 10), background="#d4edda", foreground="#333")  # Soft mint
        style.map("Excel.TButton", background=[("active", "#c3e6cb")])
        style.configure("Text.TButton", padding=6, font=('Segoe UI', 10), background="#fff3cd", foreground="#333")  # Soft yellow
        style.map("Text.TButton", background=[("active", "#ffe6b3")])
        style.configure("Generate.TButton", padding=6, font=('Segoe UI', 10), background="#e2e3f5", foreground="#333")  # Soft lilac
        style.map("Generate.TButton", background=[("active", "#d1d2f0")])
        style.configure("Sign.TButton", padding=6, font=('Segoe UI', 10), background="#c3e6cb", foreground="#333")  # Soft green
        style.map("Sign.TButton", background=[("active", "#b1d8b7")])
        style.configure("Verify.TButton", padding=6, font=('Segoe UI', 10), background="#ffddd2", foreground="#333")  # Soft coral
        style.map("Verify.TButton", background=[("active", "#ffccb3")])
        style.configure("Reset.TButton", padding=6, font=('Segoe UI', 10), background="#f8d7da", foreground="#333")  # Soft red
        style.map("Reset.TButton", background=[("active", "#f5c6cb")])
        style.configure("Save.TButton", padding=6, font=('Segoe UI', 10), background="#d1e7dd", foreground="#333")  # Soft green
        style.map("Save.TButton", background=[("active", "#c1e1d1")])
        style.configure("Upload.TButton", padding=6, font=('Segoe UI', 10), background="#fff3cd", foreground="#333")  # Soft yellow
        style.map("Upload.TButton", background=[("active", "#ffe6b3")])

    def setup_widgets(self):
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True, fill='both', padx=5, pady=5)

        # Configure grid to enforce three equal-width columns
        main_frame.grid_columnconfigure(0, weight=1, uniform="equalwidth")
        main_frame.grid_columnconfigure(1, weight=1, uniform="equalwidth")
        main_frame.grid_columnconfigure(2, weight=1, uniform="equalwidth")
        main_frame.grid_rowconfigure(0, weight=1)

        # Parameter section (left)
        self.create_param_section(main_frame, 0)
        # Sign section (middle)
        self.create_sign_section(main_frame, 1)
        # Verify section (right)
        self.create_verify_section(main_frame, 2)

    def create_param_section(self, parent, column):
        frame = ttk.LabelFrame(parent, text="Tham số DSA", padding=10, style="Param.TLabelframe")
        frame.grid(row=0, column=column, sticky="nsew", padx=5, pady=5)

        fields = [
            ("p:", self.p_var),
            ("q:", self.q_var),
            ("h:", self.h_var),
            ("g:", self.g_var),
            ("x (private key):", self.x_var),
            ("y (public key):", self.y_var)
        ]

        for i, (label, var) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky="e", padx=5, pady=2)
            ttk.Entry(frame, textvariable=var, width=20).grid(row=i, column=1, sticky="w", padx=5, pady=2)

        ttk.Button(frame, text="🎲 Sinh ngẫu nhiên tham số", command=self.generate_all_params, style="Generate.TButton").grid(row=len(fields), column=0, columnspan=2, pady=5, sticky="ew")
        ttk.Button(frame, text="🔄 Reset dữ liệu", command=self.reset_params, style="Reset.TButton").grid(row=len(fields) + 1, column=0, columnspan=2, pady=5, sticky="ew")
    def transfer_to_verify(self):
        if self.r is None or self.s is None or self.signature_hash is None or not self.message:
            messagebox.showwarning("Cảnh báo", "Vui lòng ký nội dung trước khi chuyển!")
            return

        self.verify_text_input.delete("1.0", tk.END)
        self.verify_text_input.insert(tk.END, self.message)

        self.verify_hash_entry.delete(0, tk.END)
        self.verify_hash_entry.insert(0, self.signature_hash)

        messagebox.showinfo("Thành công", "Đã chuyển nội dung và hash sang tab xác thực.")


    def create_sign_section(self, parent, column):
        frame = ttk.LabelFrame(parent, text="Ký nội dung", padding=10, style="Sign.TLabelframe")
        frame.grid(row=0, column=column, sticky="nsew", padx=5, pady=5)

        ttk.Label(frame, text="Nội dung cần ký:").pack(anchor='w', fill='x')

        btn_frame = ttk.Frame(frame, style="ButtonFrame.TFrame")
        btn_frame.pack(fill='x', pady=(0, 5))

        ttk.Button(btn_frame, text="📝 Mở file Word (.docx)", command=self.load_word_file_to_sign, style="Word.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📄 Mở file PDF (.pdf)", command=lambda: self.load_pdf_file(self.text_input), style="PDF.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📊 Mở file Excel (.xlsx)", command=lambda: self.load_excel_file(self.text_input), style="Excel.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(frame, text="➡️ Chuyển sang tab xác thực", command=self.transfer_to_verify, style="Upload.TButton").pack(pady=5, fill='x')

        self.text_input = tk.Text(frame, height=10, wrap='word', bg="#fff", font=('Segoe UI', 10))
        self.text_input.pack(fill="both", expand=True, pady=5)

        ttk.Button(frame, text="✍️ Tạo chữ ký", command=self.sign_message, style="Sign.TButton").pack(pady=5, fill='x')
        ttk.Button(frame, text="💾 Lưu chữ ký", command=self.save_signature, style="Save.TButton").pack(pady=5, fill='x')

        ttk.Label(frame, text="Kết quả (r, s, hash):").pack(anchor='w', pady=(10, 0), fill='x')
        self.signature_text = tk.Text(frame, height=5, wrap='word', state='disabled', bg="#fff", font=('Segoe UI', 10))
        self.signature_text.pack(fill="both", expand=True, pady=5)

    def create_verify_section(self, parent, column):
        frame = ttk.LabelFrame(parent, text="Xem chi tiết", padding=10, style="Verify.TLabelframe")
        frame.grid(row=0, column=column, sticky="nsew", padx=5, pady=5)

        ttk.Label(frame, text="Nội dung cần xác thực:").pack(anchor='w', fill='x')

        self.verify_text_input = tk.Text(frame, height=10, wrap='word', bg="#fff", font=('Segoe UI', 10))
        self.verify_text_input.pack(fill="both", expand=True, pady=5)

        btn_frame = ttk.Frame(frame, style="ButtonFrame.TFrame")
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text="📝 Mở file Word (.docx)", command=self.load_word_file_to_verify, style="Word.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📄 Mở file PDF (.pdf)", command=lambda: self.load_pdf_file(self.verify_text_input), style="PDF.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📊 Mở file Excel (.xlsx)", command=lambda: self.load_excel_file(self.verify_text_input), style="Excel.TButton").pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📋 Mở file Text (.txt)", command=self.load_text_file_to_verify, style="Text.TButton").pack(side='left', padx=(0, 5))

        ttk.Label(frame, text="Hash chữ ký:").pack(anchor='w', fill='x')
        self.verify_hash_entry = ttk.Entry(frame, font=('Segoe UI', 10))
        self.verify_hash_entry.pack(fill='x', pady=5)

        ttk.Button(frame, text="📑 Tải lên file chữ ký", command=self.load_signature_file, style="Upload.TButton").pack(pady=5, fill='x')
        ttk.Button(frame, text="✅ Xác thực", command=self.verify_signature, style="Verify.TButton").pack(pady=10, fill='x')
        self.result_label = tk.Label(frame, text="", font=('Segoe UI', 11, 'bold'), bg="#ffe5d9", fg="#333")
        self.result_label.pack(fill='x')

    def load_pdf_file(self, text_widget):
        file_path = filedialog.askopenfilename(
            title="Chọn file PDF (.pdf)",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if not file_path:
            return
        try:
            doc = fitz.open(file_path)
            full_text = []
            for page in doc:
                full_text.append(page.get_text())
            content = '\n'.join(full_text)
            text_widget.delete("1.0", tk.END)
            text_widget.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file PDF.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file PDF: {e}")

    def load_excel_file(self, text_widget):
        file_path = filedialog.askopenfilename(
            title="Chọn file Excel (.xlsx)",
            filetypes=[("Excel Files", "*.xlsx")]
        )
        if not file_path:
            return
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            content = []
            for row in sheet.iter_rows(values_only=True):
                content.append('\t'.join([str(cell) if cell is not None else '' for cell in row]))
            text_data = '\n'.join(content)
            text_widget.delete("1.0", "end")
            text_widget.insert("end", text_data)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Excel.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Excel: {e}")

    def show_detail_popup(self, content):
        detail_win = tk.Toplevel(self.root)
        detail_win.title("Chi tiết nội dung")
        detail_win.geometry("600x400")
        detail_win.configure(bg="#f5f5dc")
        text = tk.Text(detail_win, wrap='word', bg="#fff", font=('Segoe UI', 10))
        text.insert(tk.END, content)
        text.config(state='disabled')
        text.pack(fill='both', expand=True, padx=10, pady=10)

    def generate_all_params(self):
        self.q = random.choice([101, 103, 107])
        self.p = self.q * random.randint(70, 100) + 1
        while not is_prime(self.p):
            self.p += self.q
        self.h = random.randint(2, self.p - 2)
        self.g = calculate_g(self.p, self.q, self.h)
        self.x = random.randint(1, self.q - 1)
        self.y = pow(self.g, self.x, self.p)

        self.p_var.set(str(self.p))
        self.q_var.set(str(self.q))
        self.h_var.set(str(self.h))
        self.g_var.set(str(self.g))
        self.x_var.set(str(self.x))
        self.y_var.set(str(self.y))

    def reset_params(self):
        self.p_var.set("")
        self.q_var.set("")
        self.h_var.set("")
        self.g_var.set("")
        self.x_var.set("")
        self.y_var.set("")
        self.r = None
        self.s = None
        self.signature_hash = None

    def load_word_file_to_sign(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file Word (.docx)",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not file_path:
            return
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            self.text_input.delete("1.0", tk.END)
            self.text_input.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Word vào phần ký.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Word: {e}")

    def load_word_file_to_verify(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file Word (.docx)",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not file_path:
            return
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            self.verify_text_input.delete("1.0", tk.END)
            self.verify_text_input.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Word để xác thực.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Word: {e}")

    def load_text_file_to_verify(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file Text (.txt)",
            filetypes=[("Text Files", "*.txt")]
        )
        if not file_path:
            return
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            self.verify_text_input.delete("1.0", tk.END)
            self.verify_text_input.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Text để xác thực.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Text: {e}")

    def load_signature_file(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file chữ ký (.txt)",
            filetypes=[("Text Files", "*.txt")]
        )
        if not file_path:
            return
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                if len(lines) >= 3:
                    self.r = int(lines[0].split('=')[1].strip())
                    self.s = int(lines[1].split('=')[1].strip())
                    self.signature_hash = lines[2].split(':')[1].strip()
                    self.verify_hash_entry.delete(0, tk.END)
                    self.verify_hash_entry.insert(0, self.signature_hash)
                    messagebox.showinfo("Thành công", "Đã tải file chữ ký.")
                else:
                    raise ValueError("File chữ ký không hợp lệ!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file chữ ký: {e}")

    def save_signature(self):
        if self.r is None or self.s is None or self.signature_hash is None:
            messagebox.showwarning("Cảnh báo", "Vui lòng tạo chữ ký trước khi lưu!")
            return
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt")],
            title="Lưu chữ ký"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(f"r = {self.r}\n")
                    file.write(f"s = {self.s}\n")
                    file.write(f"Hash: {self.signature_hash}")
                messagebox.showinfo("Thành công", f"Đã lưu chữ ký vào {file_path}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể lưu file: {e}")

    def sign_message(self):
        try:
            message = self.text_input.get("1.0", tk.END).strip()
            if not message:
                messagebox.showwarning("Cảnh báo", "Vui lòng nhập nội dung!")
                return

            self.p = int(self.p_var.get())
            self.q = int(self.q_var.get())
            self.h = int(self.h_var.get())
            self.g = int(self.g_var.get())
            self.x = int(self.x_var.get())
            self.y = int(self.y_var.get())

            self.k = generate_k(self.q)
            self.r, self.s = sign(message, self.p, self.q, self.g, self.x, self.k)
            self.signature_hash = hash_signature(self.r, self.s)

            self.signature_text.config(state='normal')
            self.signature_text.delete("1.0", tk.END)
            self.signature_text.insert(tk.END, f"r = {self.r}\ns = {self.s}\nHash: {self.signature_hash}")
            self.signature_text.config(state='disabled')

            self.message = message
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def verify_signature(self):
        try:
            message = self.verify_text_input.get("1.0", tk.END).strip()
            signature_hash_input = self.verify_hash_entry.get().strip()

            if not message or not signature_hash_input:
                messagebox.showwarning("Cảnh báo", "Vui lòng tải lên văn bản và file chữ ký!")
                return

            self.p = int(self.p_var.get())
            self.q = int(self.q_var.get())
            self.h = int(self.h_var.get())
            self.g = int(self.g_var.get())
            self.x = int(self.x_var.get())
            self.y = int(self.y_var.get())

            if self.r is None or self.s is None:
                messagebox.showwarning("Cảnh báo", "Vui lòng tải lên file chữ ký hợp lệ!")
                return

            if signature_hash_input != self.signature_hash:
                self.result_label.config(text="Hash chữ ký không đúng!", fg="red")
                return

            verified = verify(message, self.p, self.q, self.g, self.y, self.r, self.s)
            if verified:
                self.result_label.config(text="Xác thực chữ ký thành công ✅", fg="green")
            else:
                self.result_label.config(text="Xác thực chữ ký thất bại ❌", fg="red")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = DSAGui(root)
    root.mainloop()